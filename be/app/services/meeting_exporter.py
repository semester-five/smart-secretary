from __future__ import annotations

import html
import unicodedata
from dataclasses import dataclass
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from app.models.project import ActionItem, Meeting, MeetingSummary, MeetingVersion, Project, Speaker, TranscriptSegment
from app.models.user import User


@dataclass(frozen=True)
class ExportContent:
    project: Project
    meeting: Meeting
    version: MeetingVersion | None
    summary: MeetingSummary | None
    speakers: list[Speaker]
    action_items: list[ActionItem]
    transcript_segments: list[TranscriptSegment]
    creator: User | None
    exported_by: User


def make_export_filename(meeting: Meeting, file_format: str) -> str:
    safe_title = _ascii_slug(meeting.title) or "meeting"
    return f"{safe_title}-minutes.{file_format}"


def build_docx(content: ExportContent) -> bytes:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    document.add_heading("Meeting Minutes", level=0)
    _add_docx_metadata_table(document, content)

    document.add_heading("Meeting Summary", level=1)
    document.add_paragraph(content.summary.summary_text if content.summary else "No summary available.")

    document.add_heading("Key Points", level=1)
    _add_docx_list(document, _extract_items(content.summary.key_points_json if content.summary else None), "No key points available.")

    document.add_heading("Decisions", level=1)
    _add_docx_list(document, _extract_items(content.summary.decisions_json if content.summary else None), "No decisions available.")

    document.add_heading("Follow-up Action Items", level=1)
    if content.action_items:
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Task", "Assignee", "Due Date", "Priority", "Status"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for item in content.action_items:
            cells = table.add_row().cells
            cells[0].text = _join_title_description(item.title, item.description)
            cells[1].text = item.assignee_text or "-"
            cells[2].text = _format_date(item.due_date)
            cells[3].text = item.priority
            cells[4].text = item.status
    else:
        document.add_paragraph("No follow-up action items available.")

    document.add_heading("Transcript", level=1)
    if content.transcript_segments:
        speaker_names = _speaker_names(content.speakers)
        for segment in content.transcript_segments:
            speaker_name = speaker_names.get(segment.speaker_id, "Speaker")
            document.add_paragraph(f"{_format_duration(segment.start_ms)} - {speaker_name}: {segment.text}")
    else:
        document.add_paragraph("No transcript available.")

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph(f"Exported at: {_format_datetime(datetime.now().astimezone())}")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf(content: ExportContent) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = _register_pdf_font()
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="MeetingTitle", fontName=font_name, fontSize=18, leading=24, alignment=TA_CENTER))
    styles.add(ParagraphStyle(name="MeetingHeading", fontName=font_name, fontSize=13, leading=18, spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle(name="MeetingBody", fontName=font_name, fontSize=9.5, leading=14))
    styles.add(ParagraphStyle(name="MeetingBullet", fontName=font_name, fontSize=9.5, leading=14, leftIndent=14))

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.6 * cm,
        leftMargin=1.6 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.4 * cm,
        title=content.meeting.title,
    )

    story: list[Any] = [
        Paragraph("Meeting Minutes", styles["MeetingTitle"]),
        Spacer(1, 12),
        _build_pdf_metadata_table(content, font_name, styles["MeetingBody"]),
        Paragraph("Meeting Summary", styles["MeetingHeading"]),
        Paragraph(_escape(content.summary.summary_text if content.summary else "No summary available."), styles["MeetingBody"]),
        Paragraph("Key Points", styles["MeetingHeading"]),
    ]
    story.extend(_pdf_bullets(_extract_items(content.summary.key_points_json if content.summary else None), styles))
    story.append(Paragraph("Decisions", styles["MeetingHeading"]))
    story.extend(_pdf_bullets(_extract_items(content.summary.decisions_json if content.summary else None), styles))

    story.append(Paragraph("Follow-up Action Items", styles["MeetingHeading"]))
    story.append(_build_pdf_action_items_table(content.action_items, font_name, styles["MeetingBody"]))

    story.append(Paragraph("Transcript", styles["MeetingHeading"]))
    if content.transcript_segments:
        speaker_names = _speaker_names(content.speakers)
        for segment in content.transcript_segments:
            speaker_name = speaker_names.get(segment.speaker_id, "Speaker")
            story.append(
                Paragraph(
                    _escape(f"{_format_duration(segment.start_ms)} - {speaker_name}: {segment.text}"),
                    styles["MeetingBody"],
                )
            )
            story.append(Spacer(1, 4))
    else:
        story.append(Paragraph("No transcript available.", styles["MeetingBody"]))

    document.build(story)
    return buffer.getvalue()


def _add_docx_metadata_table(document: Any, content: ExportContent) -> None:
    rows = _metadata_rows(content)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _add_docx_list(document: Any, items: list[str], empty_text: str) -> None:
    if not items:
        document.add_paragraph(empty_text)
        return
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _build_pdf_metadata_table(content: ExportContent, font_name: str, body_style: Any) -> Any:
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph
    from reportlab.platypus import Table, TableStyle

    data = [[Paragraph(_escape(label), body_style), Paragraph(_escape(value), body_style)] for label, value in _metadata_rows(content)]
    table = Table(data, colWidths=[4.2 * cm, 11 * cm])
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F3F4F6")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D1D5DB")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("PADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return table


def _build_pdf_action_items_table(items: list[ActionItem], font_name: str, body_style: Any) -> Any:
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph
    from reportlab.platypus import Table, TableStyle

    if not items:
        return Paragraph("No follow-up action items available.", body_style)

    data = [
        [
            Paragraph("Task", body_style),
            Paragraph("Assignee", body_style),
            Paragraph("Due Date", body_style),
            Paragraph("Priority", body_style),
            Paragraph("Status", body_style),
        ]
    ]
    for item in items:
        data.append(
            [
                Paragraph(_escape(_join_title_description(item.title, item.description)), body_style),
                Paragraph(_escape(item.assignee_text or "-"), body_style),
                Paragraph(_escape(_format_date(item.due_date)), body_style),
                Paragraph(_escape(item.priority), body_style),
                Paragraph(_escape(item.status), body_style),
            ]
        )

    table = Table(data, repeatRows=1, colWidths=[5.4 * cm, 3.1 * cm, 2.2 * cm, 2.1 * cm, 2.4 * cm])
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E5E7EB")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D1D5DB")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("PADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def _pdf_bullets(items: list[str], styles: Any) -> list[Any]:
    from reportlab.platypus import Paragraph

    if not items:
        return [Paragraph("No data available.", styles["MeetingBody"])]
    return [Paragraph(_escape(item), styles["MeetingBullet"], bulletText="-") for item in items]


def _metadata_rows(content: ExportContent) -> list[tuple[str, str]]:
    creator_name = content.creator.full_name if content.creator else content.exported_by.full_name
    participants = ", ".join(_speaker_names(content.speakers).values()) or "Not identified"
    version_status = "final" if content.version and content.version.is_final else content.meeting.status
    return [
        ("Project", content.project.name),
        ("Meeting", content.meeting.title),
        ("Meeting Date", _format_datetime(content.meeting.meeting_date)),
        ("Participants", participants),
        ("Minutes Created By", creator_name),
        ("Minutes Status", version_status),
        ("Version", str(content.version.version_no) if content.version else "latest"),
    ]


def _extract_items(value: Any) -> list[str]:
    if not value:
        return []
    if isinstance(value, list):
        return [_stringify_item(item) for item in value if _stringify_item(item)]
    if isinstance(value, dict):
        for key in ("items", "points", "key_points", "decisions"):
            nested = value.get(key)
            if isinstance(nested, list):
                return [_stringify_item(item) for item in nested if _stringify_item(item)]
        return [f"{key}: {_stringify_item(item)}" for key, item in value.items() if _stringify_item(item)]
    return [_stringify_item(value)]


def _stringify_item(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        for key in ("text", "title", "content", "decision"):
            if key in value:
                return _stringify_item(value[key])
        return "; ".join(f"{key}: {_stringify_item(item)}" for key, item in value.items() if _stringify_item(item))
    return str(value).strip()


def _speaker_names(speakers: list[Speaker]) -> dict[Any, str]:
    names: dict[Any, str] = {}
    used: set[str] = set()
    for speaker in speakers:
        name = (speaker.display_name or speaker.speaker_label).strip()
        if not name:
            continue
        if name in used:
            continue
        names[speaker.id] = name
        used.add(name)
    return names


def _join_title_description(title: str, description: str | None) -> str:
    if not description:
        return title
    return f"{title}: {description}"


def _format_datetime(value: datetime | None) -> str:
    if value is None:
        return "-"
    return value.strftime("%d/%m/%Y %H:%M")


def _format_date(value: date | None) -> str:
    if value is None:
        return "-"
    return value.strftime("%d/%m/%Y")


def _format_duration(ms: int) -> str:
    total_seconds = max(0, ms // 1000)
    minutes, seconds = divmod(total_seconds, 60)
    hours, minutes = divmod(minutes, 60)
    if hours:
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    return f"{minutes:02d}:{seconds:02d}"


def _escape(value: str) -> str:
    return html.escape(value).replace("\n", "<br/>")


def _ascii_slug(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii")
    safe_title = "".join(ch if ch.isalnum() else "-" for ch in ascii_value.lower()).strip("-")
    return "-".join(part for part in safe_title.split("-") if part)


def _register_pdf_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            font_name = candidate.stem
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, str(candidate)))
            return font_name
    return "Helvetica"
